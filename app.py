import streamlit as st
import pandas as pd
import numpy as np
from datetime import timedelta
import io
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
import math
# from streamlit_gsheets import GSheetsConnection

def main():
    hospitals = sorted(["Cardiac Research and Surgery Center", "UMC National Center for Maternal and Child Health", "Municipal Children's Hospital 2", "UMC Children's Rehabilitation Center", "Municipal Children's Hospital 3"])
    
    diseases_list = {
        'Pediatric Cardiology': {
            'Common diseases': ['Aortic valve insufficiency', 'Atrial septal defect', 'Atrioventricular septal defect', 'Bicuspid aortic valve', 'Dilated cardiomyopathy', 'Familial Hypercholesterolemia', 'Familial hypercholesterolemia, heterozygous', 'Hypertrophic cardiomyopathy', 'Hypoplasia of right ventricle', 'Idiopathic pulmonary hypertension', 'Mitral valve regurgitation', 'Paroxysmal atrioventricular nodal reentry tachycardia', 'Pulmonary trunk atresia', 'Pulmonary valve regurgitation', 'Supraventricular tachycardia', 'Tetralogy of Fallot', 'Total anomalous venous drainage of pulmonary veins', 'Tricuspid valve regurgitation', 'Truncus arteriosus', 'VSD', 'Ventricular Septal Defect', 'WPW syndrome', 'Wolff-Parkinson-White syndrome (WPW)'],
            'Frequency': [0.02631578947368421, 0.13157894736842105, 0.02631578947368421, 0.02631578947368421, 0.13157894736842105, 0.02631578947368421, 0.05263157894736842, 0.02631578947368421, 0.02631578947368421, 0.02631578947368421, 0.02631578947368421, 0.02631578947368421, 0.02631578947368421, 0.02631578947368421, 0.02631578947368421, 0.05263157894736842, 0.02631578947368421, 0.02631578947368421, 0.02631578947368421, 0.02631578947368421, 0.10526315789473684, 0.07894736842105263, 0.02631578947368421],
        },
        'Pediatric Gastroenterology': {
            'Common diseases': ["Crohn's Disease", 'Chronic hepatitis', 'Ulcerative Colitis', 'Protein–energy malnutrition', 'Behcet’s disease?', 'Celiac disease', 'Cystic Fibrosis', 'Autoimmune hepatitis', 'Dolichocolon', 'Hirshsprung disease', 'Short bowel syndrome ', 'Eosinophilic Esophagitis'],
            'Frequency': (0.36585365853658536, 0.024390243902439025, 0.2926829268292683, 0.0975609756097561, 0.024390243902439025, 0.04878048780487805, 0.024390243902439025, 0.024390243902439025, 0.024390243902439025, 0.024390243902439025, 0.024390243902439025, 0.024390243902439025),
        },
        'Pediatric Rheumatology': {
            'Common diseases': ['Juvenile Dermatomyositis', 'JIA enthesitis associated', 'JIA oligoarticular', 'JIA polyarticular', 'JIA psoriatic', 'JIA systemic', 'Juvenile Scleroderma', 'Lupus nephritis', 'MAS', 'Reactive arthritis', 'Systemic lupus erythematosus'],
            'Frequency': [0.125, 0.025, 0.25, 0.05, 0.025, 0.05, 0.2, 0.025, 0.025, 0.075, 0.15],
        },
        'Pediatric Endocrinology': {
            'Common diseases': ['Androgen insensitivity syndrome', 'Congenital Adrenal Hyperplasia', 'Diabetes mellitus, type 1', 'Diabetes mellitus, type 2', 'Diabetic ketoacidosis', 'Diabetic neuropathy', 'Growth delay', 'Hypergonadotropic hypogonadism', 'Morbid Obesity', 'Obesity stage 1', 'Osteogenesis imperfecta', 'Postoperative hypopituitarism', 'Rabson-Mendenhall syndrome', 'Short stature', 'Turner syndrome', 'X linked hereditary hypophosphatemia', 'precocious puberty'],
            'Frequency': [0.029411764705882353, 0.058823529411764705, 0.35294117647058826, 0.029411764705882353, 0.029411764705882353, 0.029411764705882353, 0.029411764705882353, 0.029411764705882353, 0.029411764705882353, 0.029411764705882353, 0.14705882352941177, 0.058823529411764705, 0.029411764705882353, 0.029411764705882353, 0.029411764705882353, 0.029411764705882353, 0.029411764705882353],
        },
        'Pediatric Radiology': {
            'us': {
                'Common diseases': ['Hepatomegaly', 'Hepatosplenomegaly', 'Splenomegaly', 'Nephromegaly', 'Post-nephrectomy', 'Pericardial effusion', 'Pleural effusion'],
                'Frequency': (0.14, 0.14, 0.14, 0.14, 0.14, 0.15, 0.15),
            },
            'xray': {
                'Common diseases': ['Pleural effusion', 'Pneumonia', 'Cardiomegaly', 'Scoliosis'],
                'Frequency': (0.25, 0.25, 0.25, 0.25),
            },
            'ct': {
                'Common diseases': ['Encephalopathy', 'Wilms tumor', 'Neuroblastoma', 'Cerebral aneurysm', 'Spinal protrusion, L5-S1', 'Spinal protrusion, L4-L5'],
                'Frequency': (0.17, 0.17, 0.17, 0.17, 0.16, 0.16),
            },
            'mri': {
                'Common diseases': ['Atherosclerotic changes in the coronary arteries', 'Ischemic stroke', 'Histiocytosis', 'Brain tumor', 'Hemorrhagic stroke'],
                'Frequency': (0.20, 0.20, 0.20, 0.20, 0.20),
            },
        },
        'Pediatric Emergencies': {
            'Common diseases': ['Acute abdomen (gerd?)', 'Ankle sprain', 'Appendicitis', 'Burns involving less than 10% of body surface', 'Clavicle fracture', 'Constipation', 'Dislocation of the elbow', 'Fall from height, soft tissue bruises', 'Foreign body aspiration', 'Functional dyspepsia', 'GERD', 'Hand burn', 'Head concussions', 'Hypoglycemia', 'Infantile colic', 'Ingrown toenail', 'Intestinal obstruction', 'Intussusception', 'Leg fracture', 'Mesenteric lymphadenitis', 'Soft tissue bruises', 'UTI', 'Ulnar fracture', 'Wrist tendon injury'],
            'Frequency': [0.03571428571428571, 0.07142857142857142, 0.07142857142857142, 0.03571428571428571, 0.03571428571428571, 0.03571428571428571, 0.03571428571428571, 0.03571428571428571, 0.07142857142857142, 0.03571428571428571, 0.03571428571428571, 0.03571428571428571, 0.03571428571428571, 0.03571428571428571, 0.03571428571428571, 0.03571428571428571, 0.03571428571428571, 0.03571428571428571, 0.03571428571428571, 0.03571428571428571, 0.03571428571428571, 0.07142857142857142, 0.03571428571428571, 0.03571428571428571],
        },
        'General Pediatrics and Pediatric Rehabilitation': {
            'Common diseases': ['cerebral palsy', 'autism', 'down syndrome'],
            'Frequency': (0.5, 0.3, 0.2),
        },
        'Pediatric Infectious Diseases': {
            'Common diseases': ['Acute Gastroenteritis', 'Acute Gastroenteritis Enterovirus', 'Acute bacterial tonsillitis', 'Acute bronchitis', 'Acute tonsillitis', 'Acute unilateral bacterial pneumonia of the right lower lobe', 'Bacterial meningitis', 'Croup', 'Measles', 'Pneumonia', 'Rhinopharyngitis', 'Scarlet fever', 'Shingles HZV infection', 'Staphylococcus infection', 'URTI', 'URTI with febrile seizures', 'VZV infection (chickenpox)', 'Viral meningitis'],
            'Frequency': [0.14705882352941177, 0.029411764705882353, 0.029411764705882353, 0.08823529411764706, 0.08823529411764706, 0.029411764705882353, 0.08823529411764706, 0.058823529411764705, 0.058823529411764705, 0.058823529411764705, 0.029411764705882353, 0.029411764705882353, 0.029411764705882353, 0.029411764705882353, 0.029411764705882353, 0.029411764705882353, 0.029411764705882353, 0.11764705882352941],
        },
        'Pediatric Oncology': {
            'Common diseases': ['Atypical teratoid rhabdoid tumor', 'Burkitt’s lymphoma', 'CNS tumour', 'Desmoplastic nodular medulloblastoma', 'Ganglioneuroblastoma', 'Hodgkin’s lymphoma', 'Medullary kidney carcinoma', 'Nephroblastoma recurrent', 'Nephroblastoma with mets', 'Neuroblastoma', 'Neuroblastoma recurrent', 'Neuroblastoma w/ mets', 'Osteogenic sarcoma', 'Retinoblastoma', 'Urinary bladder carcinoma'],
            'Frequency': [0.09523809523809523, 0.047619047619047616, 0.09523809523809523, 0.047619047619047616, 0.09523809523809523, 0.047619047619047616, 0.047619047619047616, 0.047619047619047616, 0.047619047619047616, 0.09523809523809523, 0.047619047619047616, 0.047619047619047616, 0.14285714285714285, 0.047619047619047616, 0.047619047619047616],
        },
        'Pediatric Hematology': {
            'Common diseases': ['ALL', 'AML', 'CML', 'Lymphoma'],
            'Frequency': (0.3, 0.2, 0.3, 0.2),
        },
        'Pediatric Nephrology': {
            'Common diseases': ['Post-infectious glomerulonephritis ', 'Nephrotic syndrome', 'Acute kidney injury', 'Chronic kidney disease', 'Lupus nephritis'],
            'Frequency': (0.3, 0.3, 0.1, 0.2, 0.1),
        },
        'Pediatric Intensive Care Unit': {
            'Common diseases': ['Respiratory failure', 'Acute bronchiolitis', 'Acute respiratory infection', 'Postoperative complications', 'Hemolytic uremic syndrome', 'Pneumonia', 'Sepsis'],
            'Frequency': (0.15, 0.1, 0.2, 0.1, 0.05, 0.2, 0.2),
        },
        'Neonatology': {
            'Common diseases': ['ABO Incompatibility', 'Anemia of a preterm baby', 'Asphyxia, cephalohematoma', 'Asphyxia, meconium aspiration', 'Aspiration pneumonia', 'Blood type Incompatibility', 'Bronchopulmonary Dysplasia', 'Cleft palate', 'Coarctation of aorta', 'Congenital anomaly of the urinary tract', 'Congenital pneumonia', 'Extremely low birth weight', 'Hydrocephalus', 'Hypoplasia of an aorta', 'Hypoxic-ischemic encephalopathy', 'Neonatal Jaundice', 'Neonatal Purpura Fulminans', 'Patent Ductus Arteriosus', 'Preterm baby', 'Preterm newborn', 'RDS', 'Respiratory Distress Syndrome', 'Subglottic stenosis', 'TORCH infection', 'Transient Tachypnea of the Newborn'],
            'Frequency': [0.023809523809523808, 0.07142857142857142, 0.07142857142857142, 0.023809523809523808, 0.023809523809523808, 0.023809523809523808, 0.047619047619047616, 0.023809523809523808, 0.023809523809523808, 0.023809523809523808, 0.047619047619047616, 0.023809523809523808, 0.023809523809523808, 0.023809523809523808, 0.07142857142857142, 0.14285714285714285, 0.023809523809523808, 0.023809523809523808, 0.023809523809523808, 0.023809523809523808, 0.07142857142857142, 0.047619047619047616, 0.023809523809523808, 0.023809523809523808, 0.047619047619047616],
        },
        'Pediatric Allergology, Immunology and Pulmonology': {
            'Common diseases': ['Allegic rhinitis', 'Allergy', 'Atopic dermatitis', 'Bronchial asthma', 'Immune deficiency ?', 'Immunodeficiency? CMV infection', 'Primary immune deficiency', 'Rickets', 'Urticaria'],
            'Frequency': [0.05555555555555555, 0.35185185185185186, 0.018518518518518517, 0.2037037037037037, 0.09259259259259259, 0.018518518518518517, 0.07407407407407407, 0.018518518518518517, 0.16666666666666666],
        },
        'Neonatal Intensive Care Unit': {
            'Common diseases': ['Asphyxia', 'Aspiration pneumonia', 'Blood type Incompatibility', 'Bronchopulmonary Dysplasia', 'Coarctation of aorta', 'Diaphragmatic hernia', 'Gastroschisis', 'Hemolytic disease', 'Hirschsprung disease', 'Intrauterine Growth Restriction', 'Intrauterine growth retardation ', 'Meconium aspiration ', 'NEC', 'Necrotizing enterocolitis', 'Pyloric stenosis', 'Respiratory Distress Syndrome', 'Retinopathy of prematurity', 'Sepsis'],
            'Frequency': [0.02564102564102564, 0.05128205128205128, 0.05128205128205128, 0.1282051282051282, 0.02564102564102564, 0.05128205128205128, 0.05128205128205128, 0.05128205128205128, 0.02564102564102564, 0.05128205128205128, 0.02564102564102564, 0.02564102564102564, 0.05128205128205128, 0.02564102564102564, 0.05128205128205128, 0.10256410256410256, 0.15384615384615385, 0.05128205128205128],
        },
        'Pediatric Neurology': {
            'Common diseases': ['Autistic Spectrum Disorder', 'Acute viral encephalitis', 'Cerebral Palsy', 'Focal seizures', 'Generalized Epilepsy', 'Gullain Barre Syndrome', 'Meningocele spina bifida', 'Metabolic disorder?', 'Muscular Dystrophies', 'Myotonic Dystrophy', 'Neurofibromatosis type I', 'Psychomotor delay', 'Speech delay', 'Spinal Muscular Atrophy', 'Subacute sclerosing panencephalitis', 'Tourette syndrome', "Wilson's disease"],
            'Frequency': [0.15384615384615385, 0.02564102564102564, 0.1282051282051282, 0.05128205128205128, 0.20512820512820512, 0.02564102564102564, 0.02564102564102564, 0.05128205128205128, 0.02564102564102564, 0.02564102564102564, 0.05128205128205128, 0.07692307692307693, 0.02564102564102564, 0.02564102564102564, 0.05128205128205128, 0.02564102564102564, 0.02564102564102564],
        },
    }
    
    departments = { 
        'Pediatric Cardiology': {'tutor': 'Ivanova-Razumova T.V.', 'hospital': "Cardiac Research and Surgery Center"}, 
        'Pediatric Gastroenterology': {'tutor': 'Ibrayeva A.K.', 'hospital': "UMC National Center for Maternal and Child Health"}, 
        'Pediatric Rheumatology': {'tutor': 'Assylbekova M.K.', 'hospital': "UMC National Center for Maternal and Child Health"},
        'Pediatric Endocrinology': {'tutor': 'Rakhimzhanova M.K.', 'hospital': "UMC National Center for Maternal and Child Health"},
        'Pediatric Radiology': {'tutor': 'Dautov T.B.', 'hospital': "UMC National Center for Maternal and Child Health"},
        'Pediatric Emergencies': {'tutor': 'Baigiriyev R.M.', 'hospital': "Municipal Children's Hospital 2"},
        'General Pediatrics and Pediatric Rehabilitation': {'tutor': 'Daribayev Zh.R.', 'hospital': "UMC Children's Rehabilitation Center"},
        'Pediatric Infectious Diseases': {'tutor': 'Utegenova R.B.', 'hospital': "Municipal Children's Hospital 3"},
        'Pediatric Oncology': {'tutor': 'Shaikhyzada K.K.', 'hospital': "UMC National Center for Maternal and Child Health"},
        'Pediatric Hematology': {'tutor': 'Umirbekova B.B.', 'hospital': "UMC National Center for Maternal and Child Health"},
        'Pediatric Nephrology': {'tutor': 'Rakhimzhanova S.S.', 'hospital': "UMC National Center for Maternal and Child Health"},
        'Pediatric Intensive Care Unit': {'tutor': 'Saparov A.I.', 'hospital': "UMC National Center for Maternal and Child Health"},
        'Neonatal Intensive Care Unit': {'tutor': 'Abentayeva B.A.', 'hospital': "UMC National Center for Maternal and Child Health"},
        'Pediatric Allergology, Immunology and Pulmonology': {'tutor': 'Kovzel E.F.', 'hospital': "UMC National Center for Maternal and Child Health"},
        'Neonatology': {'tutor': 'Tortayeva G.S.', 'hospital': "UMC National Center for Maternal and Child Health"},
        'Pediatric Neurology': {'tutor': 'Nauryzbayeva A.A.', 'hospital': "UMC National Center for Maternal and Child Health"},
        }

    with st.sidebar:
        name = st.text_input('Enter your first & last name:', value='Student Studentuly')
        department = st.selectbox('Choose department:', sorted(departments.keys(), key=str.lower))
        tutor = st.text_input("Enter clinical preceptor's name:", value=departments[department]['tutor'])
        hospital = st.selectbox("Choose hospital:", hospitals, index=hospitals.index(departments[department]['hospital']))
        cols = st.columns(2)
        start_date = cols[0].date_input('Enter start date', format="DD.MM.YYYY")
        end_date = cols[1].date_input('Enter end date', value=start_date+timedelta(days=53), format="DD.MM.YYYY")
        
        st.write("Select the amount of patients:")
        if department == 'Pediatric Radiology':
            patient_amount_us = st.slider("Ultrasound:", min_value=1, max_value=100, step=1, value=30)
            patient_amount_xray = st.slider("X-ray:", min_value=1, max_value=100, step=1, value=30)
            patient_amount_CT = st.slider("CT:", min_value=1, max_value=100, step=1, value=30)
            patient_amount_MRI = st.slider("MRI:", min_value=1, max_value=100, step=1, value=30)
            patient_amount = patient_amount_us + patient_amount_xray + patient_amount_CT + patient_amount_MRI
        else:    
            patient_amount = st.slider("Select the amount of patients:", min_value=1, max_value=100, step=1, value=30, label_visibility='collapsed')
        
        if department in ['Neonatal Intensive Care Unit', 'Neonatology']:
            age_range = np.arange(0, 28, 1)
        else:
            age_range = np.arange(0.1, 17, 0.1)
        if st.toggle("Preferentiate certain age group"):
            if department in ['Neonatal Intensive Care Unit', 'Neonatology']:
                st.write(f'This option is not available for the {department} department')
            else:
                preferent_age_group = st.select_slider('Preferentiable age range (years)', options=np.arange(0, 18, step=0.5), value=(0, 17.5))
                preferent_probab = st.select_slider('Percent of patients in the selected age group (%)', options=np.arange(0, 101, step=1), value=50)
                remaining_step = (((preferent_age_group[0]-0.1) + (17.9 - preferent_age_group[1]))) / (patient_amount*(100-preferent_probab)/100)
                if preferent_probab == 100:
                    age_range = np.linspace(preferent_age_group[0], preferent_age_group[1], num=patient_amount)
                else:
                    age_range = np.concatenate([np.linspace(preferent_age_group[0], preferent_age_group[1], num=round(patient_amount*preferent_probab/100)), np.arange(0.1, preferent_age_group[0], step=remaining_step), np.arange(preferent_age_group[1], 17.9, step=remaining_step)])


    # st.dataframe(diseases_list[department])

    # sheet_name = department
    # if sheet_name == 'Pediatric Radiology':
    #     sheet_name_us = 'radiology_US'
    #     sheet_name_xray = 'radiology_X_ray'
    #     sheet_name_ct = 'radiology_CT'
    #     sheet_name_mri = 'radiology_MRI'

    # # Create a connection object.
    # url = "https://docs.google.com/spreadsheets/d/1m0JxgnEUoojl_o1eRC7zRyJxF-7SxkfTS6qLUVq8Poc/edit?usp=sharing"
    # conn = st.connection("gsheets", type=GSheetsConnection)
    # df_diseases = conn.read(spreadsheet=url, usecols=[0, 1])
    
    # file = st.sidebar.file_uploader(label='Upload your file', label_visibility='collapsed')
    # if file:
    if department == 'Pediatric Radiology':
        df_diseases_us = diseases_list['Pediatric Radiology']['us']
        df_diseases_xray = diseases_list['Pediatric Radiology']['xray']
        df_diseases_ct = diseases_list['Pediatric Radiology']['ct']
        df_diseases_mri = diseases_list['Pediatric Radiology']['mri']
    else:
        df_diseases = diseases_list[department]
    
    output = {
                "№": None,
                "Age": None,
                "Gender": None,
                "Diagnosis": None,
                "Tutor's name": None,
                }

    patient_index = np.arange(1, patient_amount+1, 1)
    patients_age_numeric = np.random.choice(age_range, size=patient_amount)
    patients_age = np.empty_like(patients_age_numeric, dtype='object')
    for index, age in enumerate(patients_age_numeric):
        if department in ['Neonatology', 'Neonatal Intensive Care Unit']:
            patients_age[index] = f'{math.floor(age)} days'
        elif age >= 1:
            patients_age[index] = f'{math.floor(age)} {"years" if age > 1 else "year"}'
        else:
            patients_age[index] = f'{math.ceil(age*12)} months'

    patients_gender = np.random.choice(['male', 'female'], size=patient_amount)
    
    if department == 'Pediatric Radiology':
        diseases = np.concatenate((
            np.random.choice(df_diseases_us['Common diseases'], size=patient_amount_us, p=df_diseases_us['Frequency']),
            np.random.choice(df_diseases_xray['Common diseases'], size=patient_amount_xray, p=df_diseases_xray['Frequency']),
            np.random.choice(df_diseases_ct['Common diseases'], size=patient_amount_CT, p=df_diseases_ct['Frequency']),
            np.random.choice(df_diseases_mri['Common diseases'], size=patient_amount_MRI, p=df_diseases_mri['Frequency'])           
        ))
    else:
        diseases = np.random.choice(df_diseases['Common diseases'], size=patient_amount, p=df_diseases['Frequency'])

    if department == 'Pediatric Radiology':
        tutor_list = np.concatenate((
            ['Yurana Albayeva']*patient_amount_us,
            ['Shaddat Umbetov']*patient_amount_xray,
            ['Baurzhan Kaliev']*patient_amount_CT,
            [tutor]*patient_amount_MRI,
        ))
    else:
        tutor_list = [tutor]*patient_amount
    

    output['№'] = patient_index
    output['Age'] = patients_age
    output['Gender'] = patients_gender
    output['Diagnosis'] = diseases
    output["Tutor's name"] = tutor_list

### ------- MAIN PAGE -----------
    header = st.container()
    
    header_text = (f'''
            ***<h3 style="text-align: center;"> LOGBOOK</h3>***
            ***RESIDENT:*** {name} \n
            ***ROTATION:*** {department} \n
            ***HOSPITAL SITE:*** {hospital} \n
            ***ATTENDANCE DATE:*** **From:** {start_date}  **To:** {end_date} \n
            ***SUPERVISOR (name and surname):*** {tutor} \n
            ***SUPERVISOR (signature):*** \n
            ###
            ''')

    header.markdown(header_text, unsafe_allow_html=True)
    
    df_output = pd.DataFrame(output)

    st.data_editor(df_output, use_container_width=True)

    # Initialise the Word document
    doc = docx.Document()

    header = doc.add_paragraph("LOGBOOK")
    header.runs[0].bold = True  # Set the first run (text) to bold
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    resident = doc.add_paragraph()
    resident.add_run('RESIDENT: ').bold = True
    resident.italic = True
    resident.add_run(f'{name}')
    rotation = doc.add_paragraph()
    rotation.add_run('ROTATION: ').bold = True
    rotation.italic = True
    rotation.add_run(f'{department}')
    hospital_site = doc.add_paragraph()
    hospital_site.add_run('HOSPITAL SITE: ').bold = True
    hospital_site.italic = True
    hospital_site.add_run(f'{hospital}')
    attendance = doc.add_paragraph()
    attendance.add_run('ATTENDANCE DATE: ').bold = True
    attendance.italic = True
    attendance.add_run('From: ').bold = True
    attendance.add_run(f'{start_date} ')
    attendance.add_run('To: ').bold = True
    attendance.add_run(f'{end_date}')
    supervisor = doc.add_paragraph()
    supervisor.add_run('SUPERVISOR (name and surname): ').bold = True
    supervisor.italic = True
    supervisor.add_run(f'{tutor}')
    signature  = doc.add_paragraph()
    signature.add_run('SUPERVISOR (signature): ').bold = True
    signature.italic = True

    # Initialise the table
    t = doc.add_table(rows=(df_output.shape[0] + 1), cols=df_output.shape[1])
    # Add borders
    t.style = 'TableGrid'
    # Add the column headings
    for j in range(df_output.shape[1]):
        cell = t.cell(0, j)
        cell.text = df_output.columns[j]
        cell.paragraphs[0].runs[0].bold = True  # Set the first run (text) in the cell to bold

    # Add the body of the data frame
    for i in range(df_output.shape[0]):
        for j in range(df_output.shape[1]):
            cell = df_output.iat[i, j]
            t.cell(i + 1, j).text = str(cell)

    for cell in t.columns[0].cells:
        cell.width = Inches(0.5)
    for cell in t.columns[2].cells:
        cell.width = Inches(1)
    for cell in t.columns[3].cells:
        cell.width = Inches(2)
    for cell in t.columns[4].cells:
        cell.width = Inches(3)

    # Center-align the content in the table
    for row in t.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    bio = io.BytesIO()
    doc.save(bio)
    if doc:
        st.download_button(
            label="Click here to download",
            data=bio.getvalue(),
            file_name=f"logbook_patients_{department}.docx",
            mime="docx"
        )
# else:
#     st.subheader('Upload xls file with disease list')

if __name__ == '__main__':
    main()